from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
SCREEN_DIR = BASE_DIR / "screenshots"
OUT_DOCX = BASE_DIR / "Nota_Tecnica_Informativa_Release_Julio_Cesar_Bueno.docx"


BLUE = "2563EB"
DARK = "0F172A"
MUTED = "64748B"
LIGHT = "E8EEF5"
SOFT = "F8FAFC"
GREEN = "059669"
AMBER = "D97706"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D9E2EC", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_in: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(width * 1440)))
                tc_w.set(qn("w:type"), "dxa")


def format_run(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, size=None, bold=None, color=None, align=None, after=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        format_run(r, size=size, bold=bold, color=color)
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
        size = 15
        color = BLUE
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        size = 12.5
        color = "1F4D78"
    r = p.add_run(text)
    format_run(r, size=size, bold=True, color=color)
    return p


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        format_run(r, size=10.5, color=DARK)


def add_callout(doc, title: str, body: str, fill: str = SOFT, accent: str = BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [6.35])
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    set_cell_border(cell, "D9E2EC")
    set_cell_margins(cell, 120, 180, 120, 180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    format_run(r, size=10.5, bold=True, color=accent)
    p.paragraph_format.space_after = Pt(2)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    format_run(r2, size=10, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_key_value_table(doc, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [1.65, 4.7])
    for i, (label, value) in enumerate(rows):
        for cell in table.rows[i].cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_fill(table.cell(i, 0), LIGHT)
        p1 = table.cell(i, 0).paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label)
        format_run(r1, size=9.5, bold=True, color="1F4D78")
        p2 = table.cell(i, 1).paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        format_run(r2, size=9.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_status_table(doc, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, widths)
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        set_cell_fill(cell, LIGHT)
        set_cell_border(cell)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        format_run(r, size=9.2, bold=True, color="1F4D78")
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            set_cell_fill(cell, "FFFFFF")
            set_cell_border(cell)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            format_run(r, size=8.8, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_screenshot_pair(doc, left: tuple[str, str], right: tuple[str, str]) -> None:
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [3.05, 3.05])
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, "FFFFFF", "0")
            set_cell_margins(cell, 40, 80, 40, 80)
    for col, (filename, caption) in enumerate([left, right]):
        img_cell = table.cell(0, col)
        img_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_path = SCREEN_DIR / filename
        img_cell.paragraphs[0].add_run().add_picture(str(img_path), width=Inches(2.55))
        cap = table.cell(1, col).paragraphs[0]
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(0)
        r = cap.add_run(caption)
        format_run(r, size=8.2, color=MUTED)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Nota técnica informativa de release | Julio Cesar Bueno | Página ")
    format_run(r, size=8, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = footer.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def build() -> None:
    doc = Document()
    configure_doc(doc)
    doc.core_properties.author = "Julio Cesar Bueno"
    doc.core_properties.title = "Nota Técnica Informativa de Release - Ecossistema Acionar"
    doc.core_properties.subject = "Funcionalidades, evidências e release dos aplicativos Acionar e Patrícia Beato"

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Nota Técnica Informativa de Release")
    format_run(r, size=24, bold=True, color=DARK)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run("Ecossistema Acionar | Painel operacional + Agendamento público")
    format_run(r2, size=13.5, color="1F4D78")

    add_key_value_table(
        doc,
        [
            ("Elaborado por", "Julio Cesar Bueno"),
            ("Data", "31 de julho de 2026"),
            ("Escopo", "Dois aplicativos conectados ao mesmo Supabase/PostgreSQL"),
            ("Aplicativo 1", "Acionar - PWA operacional em GitHub Pages"),
            ("Aplicativo 2", "patriciabeato.acionar - agendamento público Lovable/TanStack"),
        ],
    )
    add_callout(
        doc,
        "Resumo executivo",
        "A release consolida um fluxo completo para captação, confirmação, execução e recebimento de atendimentos. "
        "O painel operacional concentra agenda, clientes, serviços, caixa, configurações e equipe; o app público captura "
        "solicitações de clientes com validação de horários e gravação segura via RPC no Supabase. A evolução recente reforça "
        "isolamento por profissional, horários individuais, clientes individuais e aceite exclusivo de solicitações.",
    )

    add_heading(doc, "1. Visão Geral do Ecossistema")
    add_para(
        doc,
        "Os dois aplicativos atendem papéis complementares. O Acionar é o backoffice/PWA usado por profissionais e auxiliares. "
        "O aplicativo Patrícia Beato é a experiência pública de agendamento do cliente final. Ambos compartilham o mesmo banco "
        "Supabase, mas não compartilham responsabilidades de frontend.",
        size=10.5,
        color=DARK,
    )
    add_status_table(
        doc,
        ["Componente", "Finalidade", "Usuário principal"],
        [
            ["Acionar", "Operar agenda, clientes, serviços, caixa, mensagens e equipe.", "Profissional, proprietário e auxiliar"],
            ["Patrícia Beato", "Receber solicitações públicas de agendamento com serviço, data e horário.", "Cliente final"],
            ["Supabase", "Centralizar dados, políticas, RPCs e regras de concorrência/isolamento.", "Aplicações e banco"],
        ],
        [1.55, 3.35, 1.45],
    )

    add_heading(doc, "2. Aplicativo Acionar - Funcionalidades")
    add_bullets(
        doc,
        [
            "Autenticação por Supabase Auth e fallback para profissionais/auxiliares cadastrados.",
            "Agenda com filtros por status: todos, hoje, solicitações, confirmados, em atendimento, atendidos, manutenções e cancelados.",
            "Confirmação de solicitação respeitando o profissional que aceitou; solicitações públicas deixam de aparecer aos demais após aceite.",
            "Novo agendamento manual com validação de expediente, duração do serviço, conflitos e seletor de cliente cadastrado.",
            "Cadastro de clientes isolado por profissional, com WhatsApp, edição, exclusão controlada e histórico financeiro.",
            "Cadastro de serviços, preços, duração, descrição e subserviços/opções com imagem, acréscimo e tempo adicional.",
            "Fluxo de caixa com lançamentos vinculados a atendimentos, status pago/a receber, desconto, forma e condição de pagamento.",
            "Configurações de horários por profissional, múltiplos turnos por dia, mensagens WhatsApp e endereço do estabelecimento.",
            "Notificações internas, alarme sonoro, aviso de atendimento próximo e lembretes de manutenção periódica.",
            "Suporte PWA com instalação em celular, navegação mobile e Service Worker para notificações.",
        ],
    )

    add_heading(doc, "3. Aplicativo Público Patrícia Beato - Funcionalidades")
    add_bullets(
        doc,
        [
            "Landing transacional de agendamento online, sem cadastro prévio do cliente.",
            "Coleta de nome completo e WhatsApp com máscara/validação de DDD.",
            "Listagem de serviços ativos e preços vigentes diretamente do Supabase.",
            "Suporte a subserviços/opções, com preço adicional, tempo adicional e imagem quando cadastrada.",
            "Seleção de data baseada na configuração de funcionamento publicada.",
            "Cálculo de horários livres considerando duração total e capacidade de profissionais ativos.",
            "Tratamento de conflito com sugestões de horários disponíveis.",
            "Resumo com serviço, duração, valor estimado, data, cliente e confirmação de solicitação enviada.",
            "Gravação via RPC `criar_agendamento_cliente`, mantendo status inicial `aguardando_confirmacao`.",
        ],
    )

    add_heading(doc, "4. Banco de Dados e Regras Compartilhadas")
    add_status_table(
        doc,
        ["Área", "Tabelas/Funções", "Regra relevante"],
        [
            ["Agenda", "agendamentos, listar_agendamentos_painel", "Status e visibilidade por profissional após aceite."],
            ["Clientes", "clientes", "Cliente público separado de clientes individuais por profissional."],
            ["Serviços", "servicos, tabela_precos, subservicos", "Oferta pública só mostra ativos e preços vigentes."],
            ["Equipe", "profissionais", "Auxiliares possuem agenda, clientes, caixa e horários próprios."],
            ["Caixa", "fluxo_caixa", "Recebimentos vinculados ao profissional e ao agendamento."],
            ["Segurança", "RLS, grants, RPCs", "Público acessa somente o necessário para agendar; painel usa sessão autenticada."],
        ],
        [1.3, 2.4, 2.65],
    )
    add_callout(
        doc,
        "Regra central da release",
        "Solicitações de agendamento são comuns enquanto ainda não foram aceitas. Quando um profissional ou auxiliar confirma, "
        "o agendamento passa a pertencer a esse usuário e não deve aparecer como agenda ativa dos demais.",
        fill="F0FDF4",
        accent=GREEN,
    )

    add_heading(doc, "5. Evidências de Interface")
    add_para(
        doc,
        "As capturas abaixo demonstram as telas principais. Dados exibidos em algumas telas internas são exemplos de documentação, "
        "usados apenas para evidenciar o fluxo visual.",
        size=9.7,
        color=MUTED,
    )
    add_screenshot_pair(
        doc,
        ("01_acionar_login.png", "Acionar - autenticação do painel operacional"),
        ("02_acionar_dashboard.png", "Acionar - agenda, filtros e navegação PWA"),
    )
    doc.add_page_break()
    add_heading(doc, "5. Evidências de Interface (continuação)")
    add_screenshot_pair(
        doc,
        ("03_acionar_novo_agendamento.png", "Novo agendamento manual com validação de horários"),
        ("04_acionar_selecionar_cliente.png", "Seleção de cliente cadastrado para preencher nome e WhatsApp"),
    )
    doc.add_page_break()
    add_heading(doc, "5. Evidências de Interface (continuação)")
    add_screenshot_pair(
        doc,
        ("05_acionar_clientes.png", "Gestão de clientes e base individual por profissional"),
        ("06_acionar_servicos.png", "Catálogo de serviços, preços e subserviços"),
    )
    doc.add_page_break()
    add_heading(doc, "5. Evidências de Interface (continuação)")
    add_screenshot_pair(
        doc,
        ("07_acionar_caixa.png", "Fluxo de caixa e indicadores financeiros"),
        ("08_acionar_configuracoes.png", "Horários, mensagens, alarmes e parâmetros"),
    )
    doc.add_page_break()
    add_heading(doc, "5. Evidências de Interface (continuação)")
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [4.1])
    for row in table.rows:
        cell = row.cells[0]
        set_cell_border(cell, "FFFFFF", "0")
        set_cell_margins(cell, 40, 80, 40, 80)
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 0).paragraphs[0].add_run().add_picture(
        str(SCREEN_DIR / "09_patriciabeato_agendamento_publico.png"), width=Inches(3.25)
    )
    cap = table.cell(1, 0).paragraphs[0]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("Patrícia Beato - agendamento público com serviços, preço, data e horário")
    format_run(r, size=8.5, color=MUTED)

    add_heading(doc, "6. Pontos Técnicos da Release")
    add_bullets(
        doc,
        [
            "Isolamento por `profissional_id` em agendamentos, clientes e fluxo de caixa.",
            "Backfill seguro de dados legados, preservando solicitações públicas sem profissional até o aceite.",
            "Índices para busca por profissional e unicidade de WhatsApp por contexto.",
            "RPC pública endurecida para criar agendamento com validação de nome, WhatsApp, data futura, serviço ativo e conflito de horário.",
            "Capacidade multi-profissional: o horário só é bloqueado ao atingir o número de profissionais ativos.",
            "Cache-buster do frontend atualizado para garantir entrega do JavaScript novo no GitHub Pages.",
        ],
    )

    add_heading(doc, "7. Checklist de Homologação")
    add_status_table(
        doc,
        ["Item", "Resultado esperado"],
        [
            ["Solicitação pública", "Aparece para todos até alguém confirmar."],
            ["Aceite por auxiliar", "Após confirmar, fica visível para o auxiliar que aceitou e deixa de aparecer aos demais."],
            ["Agenda manual", "Usa horários do profissional autenticado e permite selecionar cliente cadastrado."],
            ["Clientes", "Cada profissional visualiza e mantém sua própria carteira."],
            ["Caixa", "Lançamentos seguem o profissional/agendamento correspondente."],
            ["Horários", "Configuração individual por usuário, incluindo auxiliares."],
            ["App público", "Cria solicitação via RPC e exibe conflitos/sugestões quando necessário."],
        ],
        [1.9, 4.45],
    )

    add_callout(
        doc,
        "Conclusão",
        "A release deixa o produto preparado para operação com equipe: clientes, horários, caixa e agenda são individuais; "
        "solicitações públicas continuam compartilhadas até o primeiro aceite; e o cliente final permanece com um fluxo simples "
        "de agendamento online.",
        fill="FFF7ED",
        accent=AMBER,
    )

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
